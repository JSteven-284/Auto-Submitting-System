import csv
import os
from docx import Document
from docx.shared import Inches

def replace_keywords(document, keyword, replacement):
    # print(keyword)
    # print(replacement)
    for i in range(len(keyword)):
        for paragraph in document.paragraphs:
            if keyword[i] in paragraph.text:
                # print("yes")
                inline = paragraph.runs
                # Loop through each inline run
                for j in range(len(inline)):
                    # print(inline[0].text)
                    if keyword[i] in inline[j].text:
                        # print("yes")
                        inline[j].text = inline[j].text.replace(keyword[i], replacement[i])
    
def replace_pictures(document, keyword, picture_path):
    for i in range(len(keyword)):
        for paragraph in document.paragraphs:
            if keyword[i] in paragraph.text:
                # print("yes")
                inline = paragraph.runs
                # Loop through each inline run
                for j in range(len(inline)):
                    # print(inline[0].text)
                    if keyword[i] in inline[j].text:
                        # print("yes")
                        inline[j].text = inline[j].text.replace(keyword[i], '')
                        inline[j].add_picture(picture_path[i], width = Inches(5))

def generate_new_docx(template_path, csv_path, picture_path, output_path):
    document = Document(template_path)
    
    with open(csv_path, 'r', encoding='utf-8') as csv_file:
        csv_reader = csv.reader(csv_file)
        rows = []
        for row in csv_reader:
            rows.append(row)
        keyword0 = rows[0]
        # print(keyword0)
        replacement = rows[1]
        # print(keyword0)
        # print(replacement)
        # print(replacement)
        keyword1 = rows[2]
        
        replace_keywords(document, keyword0, replacement)
        replace_pictures(document, keyword1, picture_path)

    document.save(output_path)

def get_image_paths(folder_path):

    file_path1 = folder_path + "\图片一.png"
    file_path2 = folder_path + "\图片二.png"
    file_path3 = folder_path + "\图片三.png"

    image_paths = [file_path1, file_path2, file_path3]

    return image_paths

def testGenerating():
    # Specify the folder path where the images are located
    folder_path = "D:\common\Current\ASS_test\img"

    # Get the paths of all image files in the folder
    picture_path = get_image_paths(folder_path)

    # Usage
    template_path = 'template.docx'
    csv_path = 'data.csv'
    output_path = 'output.docx'
    generate_new_docx(template_path, csv_path, picture_path, output_path)
