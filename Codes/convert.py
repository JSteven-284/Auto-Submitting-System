import csv
import os
from docx import Document
from docx import Inches

def replace_keywords(document, keyword, replacement):
    # print(keyword)
    # print(replacement)
    for i in range(len(keyword)):
        for paragraph in document.paragraphs:
            if keyword[i] in paragraph.text:
                # print("yes")
                inline = paragraph.runs
                # Loop through each inline run
                for j in range(len(inline)):
                    # print(inline[0].text)
                    if keyword[i] in inline[j].text:
                        # print("yes")
                        inline[j].text = inline[j].text.replace(keyword[i], replacement[i])
    
def replace_pictures(document, keyword, picture_path):
    for i in range(len(keyword)):
        for paragraph in document.paragraphs:
            if keyword[i] in paragraph.text:
                # print("yes")
                inline = paragraph.runs
                # Loop through each inline run
                for j in range(len(inline)):
                    # print(inline[0].text)
                    if keyword[i] in inline[j].text:
                        # print("yes")
                        inline[j].text = inline[j].text.replace(keyword[i], '')
                        inline[j].add_picture(picture_path[i], width = Inches(5))

def generate_new_docx(template_path, csv_path, output_path):
    document = Document(template_path)
    
    with open(csv_path, 'r', encoding='utf-8') as csv_file:
        csv_reader = csv.reader(csv_file)
        rows = []
        for row in csv_reader:
            rows.append(row)
        keyword0 = rows[0]
        # print(keyword0)
        replacement = rows[1]
        # print(keyword0)
        # print(replacement)
        # print(replacement)
        keyword1 = rows[2]
        picture_path = ["D:\common\Current\csv2Docx\img\图片一\WIN_20230404_22_20_26_Pro.jpg",
                        "D:\common\Current\csv2Docx\img\图片二\微信图片_20230426162549.jpg",
                        "D:\common\Current\csv2Docx\img\图片三\屏幕截图 2023-05-27 120048.png"]
        
        replace_keywords(document, keyword0, replacement)
        replace_pictures(document, keyword1, picture_path)

    document.save(output_path)

def get_image_paths(folder_path):
    image_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp']  # Add more extensions if needed
    image_paths = []

    for file_name in os.listdir(folder_path):
        file_extension = os.path.splitext(file_name)[1].lower()
        if file_extension in image_extensions:
            image_path = os.path.join(folder_path, file_name)
            image_paths.append(image_path)

    return image_paths

# Specify the folder path where the images are located
folder_path = 'path/to/images'

# Get the paths of all image files in the folder
image_paths = get_image_paths(folder_path)

# Print the image paths
for path in image_paths:
    print(path)

# Usage
template_path = 'template.docx'
csv_path = 'data.csv'
output_path = 'output.docx'
generate_new_docx(template_path, csv_path, output_path)
